from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.08

sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)


def add_section_header(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): '1A1A2E',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(text)
        run.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)


def add_company_block(company, title, dates, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run(f"{company}")
    run.bold = True
    run.font.size = Pt(10.5)

    run = p.add_run(f"  |  {title}")
    run.font.size = Pt(10.5)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    run = p2.add_run(dates)
    run.font.size = Pt(9.5)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for b in bullets:
        add_bullet(b)


# ===== HEADER =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
run = p.add_run("SAQLAIN HUSSAIN")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
run = p.add_run("Senior Windows Administrator  |  L4  |  MSP Specialist")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(0)
run = p.add_run("Remote  |  saqlain.hussain@email.com  |  +91-98765-43210  |  linkedin.com/in/saqlainhussain")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ===== SUMMARY =====
add_section_header("PROFESSIONAL SUMMARY")
add_bullet(""
    "Results-driven Senior Windows Administrator with 11+ years of experience across MSP and enterprise environments. "
    "Expert in Windows Server (2012–2022), Active Directory, Microsoft 365, Azure, VMware, and automation via PowerShell. "
    "L4 escalation point for complex infrastructure incidents, driving SLA adherence, security hardening, and infrastructure modernization "
    "for 1000+ seat environments. Proven track record of leading migrations, disaster recovery planning, and zero-trust security implementations."
)

# ===== CORE SKILLS =====
add_section_header("CORE TECHNICAL SKILLS")
skills_text = (
    "Windows Server (2012–2022) | Active Directory | Group Policy | DNS | DHCP | Microsoft 365 (Exchange Online, Intune, Teams, SharePoint) | "
    "Azure (Entra ID, VMs, Networking) | AWS EC2/VPC | VMware vSphere/Hyper-V | PowerShell DSC | Packer | Terraform | Ansible | "
    "Okta SSO/MFA | Entra ID Conditional Access | Certificate Services (PKI) | ADMT | ServiceNow CMDB | "
    "SAN/NAS (NetApp, Dell EMC) | Commvault | Veeam | Cisco/Juniper Firewalls | VPN | Zero Trust Architecture | "
    "NIST/CIS Frameworks | SIEM | ITIL | Incident Management"
)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run(skills_text)
run.font.size = Pt(9.5)

# ===== EXPERIENCE =====
add_section_header("PROFESSIONAL EXPERIENCE")

# Company 15 - Most Recent (2026)
add_company_block("NexGen IT Solutions (MSP)", "Senior Windows Administrator – L4", "Jan 2026 – Present", [
    "Serve as L4 escalation point for 40+ MSP clients (3,000+ endpoints), resolving complex Windows Server, AD, and M365 incidents within SLA.",
    "Led migration of 800+ mailboxes from on-prem Exchange 2016 to Exchange Online, cutting ticketing volume by 35%.",
    "Designed and deployed Azure Virtual Desktop (AVD) environment for 200 remote users, reducing infrastructure costs by 40%.",
    "Implemented Zero Trust architecture using Entra ID Conditional Access, MFA, and Intune compliance policies across all client tenants.",
    "Automated server patching and health monitoring via PowerShell DSC and ServiceNow CMDB integration, achieving 99.8% patch compliance.",
    "Authored runbooks and SOPs; mentored L1–L3 engineers on advanced troubleshooting, AD replication, and certificate services.",
])

# Company 14 (2025-2026)
add_company_block("GlobalTech MSP", "L4 Senior Windows Systems Engineer", "Jul 2025 – Dec 2025", [
    "Provided L4 engineering support for 25+ managed clients, specializing in Active Directory forest recovery, Azure AD sync issues, and hybrid identity.",
    "Architected and executed ADMT migration for 5,000+ objects during a post-merger domain consolidation with zero data loss.",
    "Designed PKI infrastructure including certificate lifecycle automation, reducing manual cert renewal effort by 70%.",
    "Led disaster recovery drills for 15 clients, achieving RPO of 15 minutes and RTO under 4 hours using Veeam and Commvault.",
])

# Company 13 (2025)
add_company_block("EnterpriseOps Inc", "Senior Windows Infrastructure Engineer", "Jan 2025 – Jun 2025", [
    "Managed 500+ Windows Server 2022 VMs across VMware vSphere and Hyper-V for a Fortune 1000 enterprise environment.",
    "Architected golden image pipeline using Packer and PowerShell DSC for Windows Server 2022, accelerating provisioning from 3 days to 4 hours.",
    "Consolidated 12 domain controllers across 4 sites into a streamlined AD topology, reducing replication latency by 60%.",
    "Implemented SIEM integration for Windows Event Log forwarding to Splunk SOC, enhancing threat detection coverage.",
])

# Company 12 (2024-2025)
add_company_block("ApexLogic (MSP)", "L4 Systems Engineer – Windows & Azure", "Jun 2024 – Dec 2024", [
    "Acted as internal L4 escalation for complex hybrid Azure AD/Entra ID join and sync issues across 30+ client tenants.",
    "Designed and deployed Intune autopilot and compliance policies for 2,000+ Windows endpoints, reducing imaging time by 80%.",
    "Migrated 50+ legacy application servers from on-prem to Azure IaaS, utilizing Azure Migrate and ASR for minimal downtime.",
    "Built PowerShell-based automated health reporting system for client AD, DNS, and DHCP infrastructure, improving proactive issue detection by 50%.",
])

# Company 11 (2024)
add_company_block("TierPoint IT (MSP)", "Windows Infrastructure Lead – L3/L4", "Jan 2024 – May 2024", [
    "Led 10-person infrastructure team across 20 MSP clients, establishing best practices for AD, GPO, DNS, and DHCP management.",
    "Performed AD security assessment and remediation for 8 clients, resolving dormant privilege escalations and legacy trust issues.",
    "Implemented NIST SP 800-53 controls across client server environments, achieving 100% pass rate in external audit.",
    "Reduced critical incident resolution time by 45% through creation of Tier 3 runbooks and automated escalation workflows.",
])

# Company 10 (2023-2024)
add_company_block("CloudSync Global (MSP)", "Senior Windows Systems Engineer", "Mar 2023 – Dec 2023", [
    "Managed hybrid Exchange 2019 / Exchange Online environment for 15 multi-national clients with 10,000+ mailboxes.",
    "Designed and deployed Azure File Sync and FSLogix profile containers for 500+ remote users in VDI environment.",
    "Automated Windows server hardening against CIS benchmarks using PowerShell DSC, reducing audit findings by 75%.",
    "Led OKTA SSO integration for 10+ client applications, streamlining authentication and enabling adaptive MFA policies.",
])

# Company 9 (2022-2023)
add_company_block("VirtuCore IT Services (MSP)", "Senior Systems Administrator – L3", "May 2022 – Feb 2023", [
    "L3 escalation engineer for 35+ MSP clients across Windows Server, AD, M365, and networking infrastructure.",
    "Architected client onboarding process for new MSP acquisitions, standardizing AD structure, GPOs, and M365 tenant configuration.",
    "Performed domain migration and consolidation for 4 merged entities, unifying 3,000+ users under single forest.",
    "Deployed WSUS/SCCM-based patch management across 2,000+ servers, achieving 95% patch compliance within 72 hours.",
])

# Company 8 (2021-2022)
add_company_block("EliteTech Solutions (MSP)", "L3/L4 Windows Engineer", "Aug 2021 – Apr 2022", [
    "Provided Tier 3/4 support for 50+ client environments, resolving escalated AD replication, DFS, and Exchange issues.",
    "Migrated on-prem file servers to SharePoint Online and OneDrive for Business across 10 clients, managing 5 TB of data.",
    "Implemented conditional access policies and MFA via Azure AD P2, reducing phishing compromise incidents by 90%.",
    "Designed and documented DR/BCP plans for 12 clients leveraging Azure Site Recovery and Veeam replication.",
])

# Company 7 (2020-2021)
add_company_block("NetForge Managed Services (MSP)", "Lead Windows Administrator – L3", "Sep 2020 – Jul 2021", [
    "Managed Windows Server 2016/2019 fleet across 30+ client environments including AD, DNS, DHCP, and file/print services.",
    "Developed PowerShell automation scripts for user provisioning, password resets, and group membership audits — saving 20 engineer-hours per week.",
    "Led migration of 3,000+ user accounts from on-prem AD to Azure AD hybrid identity with password hash sync and seamless SSO.",
    "Configured and managed VMware vSphere clusters for client hosts, including VM provisioning, resource pools, and DRS rules.",
])

# Company 6 (2019-2020)
add_company_block("SecureLink IT (MSP)", "Senior Systems Administrator", "Nov 2019 – Aug 2020", [
    "Delivered L2/L3 support for 25+ MSP clients, specializing in Windows Server troubleshooting, M365 administration, and security remediation.",
    "Performed security hardening across 500+ servers applying CIS benchmarks and STIG requirements via Group Policy.",
    "Managed client migrations from on-prem Skype for Business to Microsoft Teams, coordinating user cutover with zero downtime.",
    "Deployed RMM agents and monitoring dashboards for proactive infrastructure health tracking and alerting.",
])

# Company 5 (2018-2019)
add_company_block("DataBridge Technologies", "Windows System Administrator", "Jun 2018 – Oct 2019", [
    "Administered 300+ Windows Server 2012/2016 environment including Active Directory, Exchange 2013, and file services.",
    "Automated routine maintenance tasks (log rotation, backup verification, user audits) using PowerShell scripts.",
    "Migrated on-premises infrastructure to AWS EC2/S3, supporting 150+ application servers with minimal business disruption.",
    "Maintained 99.9% uptime for critical systems through proactive monitoring using SolarWinds and Nagios.",
])

# Company 4 (2017-2018)
add_company_block("InfraCore Services (MSP)", "Windows Administrator – L2/L3", "Mar 2017 – May 2018", [
    "Provided L2/L3 support across 20+ client environments, resolving escalated AD, Exchange, and connectivity issues.",
    "Deployed and managed Microsoft 365 tenants for 15 small-to-medium businesses via the partner portal.",
    "Configured and maintained VMware ESXi hosts, virtual networking, and storage (iSCSI, NFS, Fibre Channel SAN).",
    "Created technical documentation and knowledge base articles for common incident resolution procedures.",
])

# Company 3 (2016-2017)
add_company_block("CloudPeak Systems (MSP)", "Systems Administrator – L2", "Aug 2016 – Feb 2017", [
    "Supported 15+ MSP clients with Windows Server, Active Directory, and desktop infrastructure issues.",
    "Managed user identity lifecycle (provisioning, deprovisioning, access reviews) across multiple AD domains.",
    "Assisted in migration project for 500+ users from Exchange 2010 to Exchange 2016 with zero data loss.",
    "Configured firewall rules, VLANs, and site-to-site VPNs for client branch office connectivity.",
])

# Company 2 (2015-2016)
add_company_block("TechNova Solutions (MSP)", "Junior Windows Administrator – L1/L2", "Mar 2015 – Jul 2016", [
    "Provided L1/L2 helpdesk support for 1,000+ users across managed client environments (Windows, M365, basic AD).",
    "Created and maintained Active Directory user accounts, security groups, and distribution lists per client policies.",
    "Assisted with patch management deployment via WSUS and third-party patching tools across 500+ endpoints.",
    "Documented troubleshooting procedures and contributed to the internal knowledge base.",
])

# Company 1 (2015)
add_company_block("NetConnect IT Services", "IT Support Engineer – L1", "Jan 2015 – Feb 2015", [
    "Provided Level 1 technical support for end-user desktops, printers, and basic network connectivity issues.",
    "Assisted senior engineers with Active Directory user management and password resets.",
    "Imaged and deployed Windows 7/8 workstations using automated deployment tools.",
    "Tracked and resolved 40+ tickets per week via ServiceNow, maintaining 90%+ first-response SLA.",
])

# ===== CERTIFICATIONS =====
add_section_header("CERTIFICATIONS")
certs = [
    "Microsoft Certified: Azure Solutions Architect Expert – 2024",
    "Microsoft Certified: Windows Server Hybrid Administrator Associate – 2023",
    "Microsoft 365 Certified: Enterprise Administrator Expert – 2023",
    "CompTIA Security+ (SY0-601) – 2022",
    "ITIL v4 Foundation – 2021",
    "VMware Certified Professional – VCP-DCV 2022",
    "Okta Certified Professional – 2023",
]
for c in certs:
    add_bullet(c)

# ===== EDUCATION =====
add_section_header("EDUCATION")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run("Bachelor of Technology in Computer Science")
run.bold = True
run.font.size = Pt(10.5)
p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(0)
p2.paragraph_format.space_before = Pt(0)
run = p2.add_run("University of Mumbai  |  2010 – 2014")
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.italic = True

# Save
output_path = "/home/sultantalib/hubgroup/Saqlain_Hussain_Resume_Windows_Admin_L4.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
